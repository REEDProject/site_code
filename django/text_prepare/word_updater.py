import io
import os
import re
import shlex
import subprocess
import tempfile

import docx

from .exceptions import TextPrepareDocumentUpdateError


# Change the format of closing @-codes from '@x \' to '@x/'. The
# original whitespace can cause problems in the parsing, but the
# editors are used to that form.
codes = [
    'a', 'ab', 'b', 'c', 'cl', 'cn', 'cnx', 'cor', 'cr', 'cym', 'deu', 'e',
    'en', 'eng', 'ex', 'f', 'fra', 'g', 'gla', 'gmh', 'gml', 'grc', 'i',
    'ita', 'j', 'k', 'l', 'lat', 'li', 'm', 'p', 'pc', 'por', 'q', 'r',
    's', 'sc', 'sd', 'sh', 'sl', 'sn', 'snc', 'snr', 'spa', 'sr', 'ss',
    'st', 'tr', 'ul', 'wlm', 'x', 'xc', 'xno']
convert_code = re.compile(r'@(({}))\s\\'.format(')|('.join(codes)))


def convert_to_docx(doc_path):
    with tempfile.TemporaryDirectory() as env_fh:
        command = '''soffice -env:UserInstallation=file://{} --headless
                     --convert-to docx {}'''.format(env_fh, doc_path)
        try:
            subprocess.check_call(shlex.split(command), cwd=os.path.dirname(
                doc_path))
        except subprocess.CalledProcessError as e:
            msg = 'Failed to convert Word document to docx format: {}'
            raise TextPrepareDocumentUpdateError(msg.format(e.output))
    return os.path.splitext(doc_path)[0] + '.docx'


def update_word(doc_path):
    try:
        doc = docx.Document(docx=doc_path)
    except ValueError:
        # doc_file may point to a non-docx file, in which case try to
        # convert it.
        docx_path = convert_to_docx(doc_path)
        doc = docx.Document(docx=docx_path)
        os.remove(docx_path)
    for paragraph in doc.paragraphs:
        paragraph.text = convert_code.sub(r'@\1/', paragraph.text)
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()
