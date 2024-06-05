"""Module containing the Document class."""

import io
import os.path
import re
import shlex
import subprocess
import tempfile
import textwrap
import unicodedata
import zipfile

import docx
from lxml import etree
import pyparsing as pp

from .document_parser import DocumentParser
from .exceptions import (TextPrepareDocumentError,
                         TextPrepareDocumentValidationError)


# Constants for the number of lines of context to display before and
# after the line on which a validation error occurs.
CONTEXT_LINES_BEFORE = 1
CONTEXT_LINES_AFTER = 2

# @-codes that need to have their format modified from "@x \" to "@x/".
CODES = [
    'a', 'ab', 'ang', 'b', 'c', 'cl', 'cnx', 'cor', 'ct', 'cym', 'en', 'eng',
    'ex', 'f', 'fra', 'g', 'ger', 'gla', 'gmh', 'gml', 'grc', 'i', 'it',
    'ita', 'l', 'lat', 'lg', 'li', 'ln', 'lni', 'm', 'md', 'p', 'pc',
    'pd', 'por', 'r', 's', 'sc', 'sd', 'sh', 'sl', 'sm', 'smc', 'smr',
    'sn', 'snc', 'snr', 'spa', 'sr', 'ss', 'st', 'tr', 'ul', 'wlm',
    'x', 'xc', 'xno']

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
XSLT_DIR = os.path.join(BASE_DIR, 'xslt')
AB_TO_P_XSLT_PATH = os.path.join(XSLT_DIR, 'ab_to_p.xsl')
ADD_AB_XSLT_PATH = os.path.join(XSLT_DIR, 'add_ab.xsl')
ADD_HEADER_XSLT_PATH = os.path.join(XSLT_DIR, 'add_header.xsl')
ADD_ID_XSLT_PATH = os.path.join(XSLT_DIR, 'add_id.xsl')
MASSAGE_FOOTNOTE_XSLT_PATH = os.path.join(XSLT_DIR, 'massage_footnote.xsl')
REMOVE_AB_XSLT_PATH = os.path.join(XSLT_DIR, 'remove_ab.xsl')
SANITISE_WORD_XSLT_PATH = os.path.join(XSLT_DIR, 'sanitise_word.xsl')
SORT_RECORDS_XSLT_PATH = os.path.join(XSLT_DIR, 'sort_records.xsl')
TIDY_BIBLS_XSLT_PATH = os.path.join(XSLT_DIR, 'tidy_bibls.xsl')

BIBL_SKELETON = '''<listBibl xmlns="http://www.tei-c.org/ns/1.0">
{}
</listBibl>'''
RECORDS_SKELETON = '''<TEI xmlns="http://www.tei-c.org/ns/1.0" xml:base="tei/records/" xml:lang="eng">
  <text>
    <group>{}</group>
  </text>
</TEI>'''


class Document:

    """Represents a complete records TEI XML document combined with a
    snippet of document descriptions, and provides methods to validate
    and convert one or more Word documents into XML."""

    def __init__(self, base_id=''):
        self._base_id = base_id
        self._doc_descs = []
        self._records = []

    def convert(self, word_file_path, line_length):
        """Converts the content of the supplied Word document into TEI XML.

        :param word_file_path: path to Word file to be converted
        :type file_path: `str`
        :param line_length: number of characters to wrap lines to
        :type line_length: `int`

        """
        text = self._get_text(word_file_path, line_length)
        results = self._validate(text)
        self._doc_descs.append(''.join(results.doc_desc))
        self._records.append(''.join(results.record))

    def _convert_at_codes(self, text):
        """Returns `text` with @-code markup converted into the form required
        by the grammar.

        :param text: text to convert
        :type text: `str`
        :rtype: `str`

        """
        # Change the format of closing @-codes from '@x \' to '@x/'. The
        # original whitespace can cause problems in the parsing, but the
        # editors are used to that form.
        for code in CODES:
            text = text.replace('@{} \\'.format(code), '@{}/'.format(code))
        return text

    def _convert_to_docx(self, doc_path):
        with tempfile.TemporaryDirectory() as env_fh:
            command = '''soffice -env:UserInstallation=file://{} --headless
                         --convert-to docx {}'''.format(env_fh, doc_path)
            try:
                subprocess.check_call(shlex.split(command),
                                      cwd=os.path.dirname(doc_path))
            except subprocess.CalledProcessError as e:
                msg = 'Failed to convert Word document to docx format: {}'
                raise TextPrepareDocumentError(msg.format(e.output))
        return os.path.splitext(doc_path)[0] + '.docx'

    def generate(self):
        """Returns the binary content of a ZIP file containing the TEI XML
        records file and a TEI XML fragment containing document
        descriptions."""
        bibls = self._generate_bibls()
        records = self._generate_tei()
        archive = io.BytesIO()
        bibl_filename = 'taxonomy-excerpt.xml'
        records_filename = '{}.xml'.format(self._base_id)
        with zipfile.ZipFile(archive, 'w') as zip_file:
            zip_file.writestr(bibl_filename, bibls.encode('utf-8'))
            zip_file.writestr(records_filename, records.encode('utf-8'))
        return archive.getvalue()

    def _generate_bibls(self):
        """Returns the TEI XML fragment containing document descriptions."""
        text = BIBL_SKELETON.format('\n'.join(self._doc_descs))
        tree = etree.ElementTree(etree.fromstring(text))
        tree = self._transform(tree, TIDY_BIBLS_XSLT_PATH)
        return etree.tostring(tree, encoding='unicode', pretty_print=True)

    def _generate_tei(self):
        """Returns the TEI XML records text."""
        text = RECORDS_SKELETON.format(''.join(self._records))
        text = unicodedata.normalize('NFC', text)
        text = self._postprocess_text(text)
        # Make some cosmetic changes that are too hard to do with XSLT
        # 1. This is rather dirty!
        text = re.sub(r'(<ab[^>]*>)[ \t\n\r\f\v]+', r'\1', text)
        text = re.sub(r'[ \t\n\r\f\v]+(</ab>)', r'\1', text)
        return text

    def _get_text(self, file_path, line_length):
        """Returns the plain text conversion of the Word file at `file_path`.

        :param file_path: path to file to extract text from
        :type file_path: `str`
        :param line_length: number of characters to wrap lines to
        :type line_length: `int`
        :rtype: `str`

        """
        text = []
        try:
            docx.Document(docx=file_path)
            docx_path = file_path
        except ValueError:
            # doc_file may point to a non-docx file, in which case try to
            # convert it.
            docx_path = self._convert_to_docx(file_path)
        # Having ensured that we have a docx file, we must first
        # manipulate one of its constituent files before opening it
        # via the third party library, to cope with the fact that
        # non-breaking hyphens are not treated as text.
        self._sanitise_word(docx_path)
        doc = docx.Document(docx=docx_path)
        os.remove(docx_path)
        for paragraph in doc.paragraphs:
            text.append(self._update_text(paragraph.text))
        return self._wrap_text('\n'.join(text), line_length)

    def _postprocess_text(self, text):
        """Returns `text` modified by applying various XSLT transformations to
        it.

        :param text: TEI XML to post-process
        :type text: `str`
        :rtype: `str`

        """
        tree = etree.ElementTree(etree.fromstring(text))
        tree = self._transform(
            tree, ADD_AB_XSLT_PATH, ADD_ID_XSLT_PATH, ADD_HEADER_XSLT_PATH,
            MASSAGE_FOOTNOTE_XSLT_PATH, REMOVE_AB_XSLT_PATH,
            AB_TO_P_XSLT_PATH, SORT_RECORDS_XSLT_PATH)
        return etree.tostring(tree, encoding='unicode', pretty_print=True)

    def _replace_word_chars(self, text):
        """Returns `text` with certain characters replaced with their
        preferred alternates."""
        text = text.replace(' ', ' ')
        text = text.replace('‘', "'")
        text = text.replace('’', "'")
        text = text.replace('“', '"')
        text = text.replace('”', '"')
        return text

    def _sanitise_word(self, docx_path):
        """Returns a sanitised form of the XML of the text content of the Word
        file at `docx_path`."""
        content_path = 'word/document.xml'
        tmp_path = docx_path + '.new'
        transform = etree.XSLT(etree.parse(SANITISE_WORD_XSLT_PATH))
        with zipfile.ZipFile(docx_path, 'r') as zip_in_file:
            with zipfile.ZipFile(tmp_path, 'w') as zip_out_file:
                for document in zip_in_file.infolist():
                    with zip_in_file.open(document) as document_file:
                        if document.filename == content_path:
                            tree = etree.parse(document_file)
                            tree = transform(tree)
                            content = etree.tostring(tree, encoding='utf-8',
                                                     pretty_print=False)
                        else:
                            content = document_file.read()
                    zip_out_file.writestr(document, content)
        os.remove(docx_path)
        os.rename(tmp_path, docx_path)

    def _transform(self, tree, *xslt_paths):
        for path in xslt_paths:
            transform = etree.XSLT(etree.parse(path))
            tree = transform(tree, base_id="'{}'".format(self._base_id))
        return tree

    def _update_text(self, text):
        text = self._replace_word_chars(text)
        text = self._convert_at_codes(text)
        # To avoid problems with Python's .format(), double all curly braces.
        text = text.replace('{', '{{').replace('}', '}}')
        text = text.replace('[[', '@DOUBLE_SQUARE_OPEN').replace(
            ']]', '@DOUBLE_SQUARE_CLOSE')
        return text

    def validate(self, word_file_path, line_length):
        """Raises an exception if this document does not conform to the @-code
        grammar."""
        text = self._get_text(word_file_path, line_length)
        self._validate(text)

    def _validate(self, text):
        """Returns `text` converted into TEI XML according to the rules
        attached to the @-code grammar. This conversion raises an
        exception if `text` does not conform to the grammar.

        :param text: textual content to be converted
        :type text: `str`
        :rtype: `str`

        """
        parser = DocumentParser()
        try:
            results = parser.parse(text)
        except (pp.ParseException, pp.ParseSyntaxException,
                pp.ParseFatalException) as e:
            # Generate the context of the error, with a pointer to the
            # column identified as its place.
            col = e.column
            line = e.line
            line_number = e.lineno
            line_index = e.lineno - 1
            split_text = text.splitlines()
            lines = []
            lines.append('Line: {}'.format(line_number))
            lines.append('Message: {}\n'.format(e))
            if line_number > CONTEXT_LINES_BEFORE:
                lines.extend(split_text[
                    line_index-CONTEXT_LINES_BEFORE:line_index])
            lines.append(line)
            lines.append(' ' * col + '^')
            try:
                lines.extend(split_text[
                    line_index+1:line_index+1+CONTEXT_LINES_AFTER])
            except:
                pass
            message = '\n'.join(lines)
            # Undo doubling of curly braces.
            message = message.replace('{{', '{').replace('}}', '}')
            raise TextPrepareDocumentValidationError(message)
        return results

    def _wrap_text(self, text, line_length):
        """Returns `text` wrapped to the specified line `length`.

        :param text: text to wrap
        :type text: `str`
        :param line_length: length of line to wrap `text` to
        :type line_length: `int`
        :rtype: `str`

        """
        # Due to slight differences in how Word and textwrap wrap
        # lines, it can happen that a closing @-code (eg, "@f \") is
        # split over two lines on the whitespace, which breaks the
        # parsing of this element. Therefore hack around this.
        text = text.replace(' \\', '豈\\')
        wrapper = textwrap.TextWrapper(width=line_length)
        text = '\n'.join([wrapper.fill(line) for line in text.splitlines()])
        text = text.replace('豈\\', ' \\')
        return text
